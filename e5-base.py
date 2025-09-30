# -*- coding: utf-8 -*-
"""
R&D 계획서 자동 생성 파이프라인 (검색 RAG 활성화 버전)
- 가이드라인/참고 JSON 로드 → 청크/임베딩/FAISS 인덱스 → Top-K(+MMR) 검색 컨텍스트 → 생성 LLM → DOCX 저장
- 메모리/안정성 기본값으로 설정 (OOM 방지 팁은 주석 참조)
"""

import os, json, random, math, re
import numpy as np
import torch
from tqdm.auto import tqdm
from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from sentence_transformers import SentenceTransformer
import faiss
from datasets import Dataset
from torch.utils.data import DataLoader
from typing import List, Dict, Any

# ==============================================================================
# 1) 환경/경로/모델 파라미터
# ==============================================================================

# --- 데이터 소스 경로 ---
GUIDELINE_FILE = "/home/alpaco/autosry/rnd_guideline.json"
RAG_JSON_FILES = ["/home/alpaco/autosry/rag_chunks500_50.json"]

# --- 모델 이름 ---
E5_NAME = "intfloat/e5-base-v2"    # 검색용 임베딩 모델(E5-base)
GEN_NAME = "skt/A.X-4.0-Light"     # 생성용 LLM
# The above code snippet defines various parameters related to device settings, search/RAG parameters,
# and a commented-out loop for reranking hits. Here is a breakdown of what the code is doing:

# --- 디바이스 ---
# torch.device(DEVICE)를 명시적으로 사용
DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
# CUDA_ID는 device_map="auto"를 사용할 경우 자동 관리되므로 주석 처리
# CUDA_ID = 1 if torch.cuda.is_available() else -1 

# --- 검색/RAG 파라미터 ---
DOCUMENT_PREFIX = "passage: "
QUERY_PREFIX    = "query: "
MAXLEN_DOC      = 512          # 문서(청크) 토큰 길이 상한
MAXLEN_QUERY    = 256          # 쿼리 토큰 길이 상한
EMBED_BATCH     = 32           # 임베딩 배치 (메모리 부족 시 8~16 추천)
TOPK= 4                # 최종 프롬프트에 넣을 스니펫 K
SEARCH_POOL     = 32           # 1차 검색 풀(N) → 여기서 MMR로 K개 선별
MMR_LAMBDA      = 0.5          # MMR trade-off (0~1), 0.5 권장


# reranked_hits = [] 
# for idx, score in hits_with_indices:
#     reranked_hits.append((idx, final_score))


# --- 생성 파라미터(메모리 안전 기본값) ---
GEN_MODE            = "pipeline_batch"  # "pipeline_batch" | "datasets_manual"
GEN_MAX_NEW_TOKENS = 1024               # 5000은 OOM 위험, 512~1024 권장
GEN_DO_SAMPLE      = False              # False면 greedy/beam; True면 temperature 등 사용
GEN_TEMPERATURE    = None               # 샘플링 사용 시만 의미

# ==============================================================================
# 2) 유틸리티
# ==============================================================================

def l2_normalize(t: torch.Tensor, dim: int = 1) -> torch.Tensor:
    """(미사용) 텐서 L2 정규화."""
    return torch.nn.functional.normalize(t, p=2, dim=dim)

def first_n_chars(text: str, n_chars: int = 350) -> str:
    """텍스트를 공백 하나로 정규화하고 앞 n_chars만 반환(프롬프트 길이 보호)."""
    t = " ".join(str(text).split())
    return t[:n_chars]

def clean_generated_text(text: str) -> str:
    """생성 결과에서 프롬프트 헤더/마크다운 기호 등을 제거해 본문만 정리."""
    output_marker = "#=========== 출력"
    # 1) 헤더 이후만 취함
    if output_marker in text:
        text = text.split(output_marker, 1)[-1].strip()
    # 2) 마크다운 볼드/헤딩 제거
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'#+\s*', '', text)
    text = text.replace('**', '')
    return text.strip()

# ==============================================================================
# 3) 데이터 로드 (가이드라인 + RAG 텍스트 통합)
# ==============================================================================

try:
    with open(GUIDELINE_FILE, "r", encoding="utf-8") as f:
        guidelines = json.load(f)  # (프롬프트에서 파일 경로로만 언급, 직접 사용 X)
except FileNotFoundError:
    print(f"[WARN] Guideline file not found: {GUIDELINE_FILE}")
    guidelines = {}
except json.JSONDecodeError:
    print(f"[WARN] Guideline file decode error: {GUIDELINE_FILE}")
    guidelines = {}

combined_text = ""  # 모든 RAG 텍스트를 이어붙여 컨텍스트 원천으로 사용
for jf in RAG_JSON_FILES:
    try:
        with open(jf, "r", encoding="utf-8") as f:
            parsed = json.load(f)
            if isinstance(parsed, dict) and "text" in parsed:
                combined_text += str(parsed["text"]) + "\n\n"
            else:
                combined_text += str(parsed) + "\n\n"
    except FileNotFoundError:
        print(f"[WARN] RAG JSON file not found: {jf}")
    except json.JSONDecodeError:
        print(f"[WARN] RAG JSON file decode error: {jf}")

print("[INFO] RAG chunks loaded. (combined_text prepared)")

# ==============================================================================
# 3.5) 청크화 + E5 임베딩 + FAISS 인덱스
# ==============================================================================

def simple_chunker(raw_text: str, target_chars: int = 600, overlap: int = 100) -> List[str]:
    """
    매우 단순한 글자수 기준 슬라이싱 청크러(+오버랩).
    - 문장/문단 단위 청크러로 개선 가능(규칙 기반/SegTok 등).
    """
    text = " ".join(str(raw_text).split())
    chunks, start, n = [], 0, len(text)
    while start < n:
        end = min(start + target_chars, n)
        chunk = text[start:end].strip()
        if len(chunk) >= 50:  # 너무 짧은 꼬리 제거
            chunks.append(chunk)
        if end == n:
            break
        start = max(0, end - overlap)  # 오버랩 유지
    return chunks

# 1) 청크 목록 구성
rag_docs: List[Dict[str, Any]] = []
if combined_text.strip():
    for i, c in enumerate(simple_chunker(combined_text, target_chars=600, overlap=100)):
        rag_docs.append({"id": i, "text": c})

print(f"[INFO] Built {len(rag_docs)} chunks for retrieval.")

# 2) E5 임베딩 모델 로드 (DEVICE 변수 사용)
e5 = SentenceTransformer(E5_NAME, device=DEVICE)
e5.max_seq_length = 512  # 안전 상한

def embed_texts_e5(texts: List[str], is_query: bool = False, batch_size: int = EMBED_BATCH) -> np.ndarray:
    """
    E5 임베딩 계산(with prefix). normalize_embeddings=True → 코사인=내적.
    """
    prefix = QUERY_PREFIX if is_query else DOCUMENT_PREFIX
    # 문자 기준 슬라이스(토큰 길이 엄밀 제한 대신 안전빵으로 과도 입력 방지)
    inputs = [ (prefix + t)[: (MAXLEN_QUERY if is_query else MAXLEN_DOC) * 4] for t in texts ]
    emb = e5.encode(inputs, batch_size=batch_size, normalize_embeddings=True, show_progress_bar=False)
    return emb.astype("float16")

# 3) 문서 임베딩 & FAISS 인덱스(IP: 코사인 등가)
if len(rag_docs) == 0:
    faiss_index = None
    doc_texts   = []
    doc_emb     = None
else:
    doc_texts   = [d["text"] for d in rag_docs]
    doc_emb     = embed_texts_e5(doc_texts, is_query=False)
    dim         = doc_emb.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(doc_emb)
    print(f"[INFO] FAISS index built. dim={dim}, n={faiss_index.ntotal}")

def mmr_select(q: np.ndarray, docs: np.ndarray, k: int, lam: float = 0.5) -> List[int]:
    """
    MMR(Maximal Marginal Relevance)로 다양성 리랭킹.
    - q: (d,), docs: (n, d) 모두 L2 정규화된 벡터(내적=코사인).
    - lam: 관련성(쿼리-문서) vs 다양성(문서-문서)의 가중.
    """
    n = docs.shape[0]
    if n == 0:
        return []
    k = min(k, n)

    # 쿼리-문서 유사도
    sim_q = docs @ q  # (n,)

    selected = []
    candidates = set(range(n))

    # 1개는 관련성 최대를 우선 선택
    first = int(np.argmax(sim_q))
    selected.append(first)
    candidates.remove(first)

    # 문서-문서 유사도(선택집합 대비 최대값) 활용
    docdoc = docs @ docs.T  # (n, n)

    while len(selected) < k and len(candidates) > 0:
        best_i, best_score = None, -1e9
        for i in candidates:
            # MMR score = lam*sim(q, i) - (1-lam)*max_j sim(i, j) for j in selected
            redundancy = np.max(docdoc[i, selected])
            score = lam * sim_q[i] - (1 - lam) * redundancy
            if score > best_score:
                best_score = score
                best_i = i
        selected.append(best_i)
        candidates.remove(best_i)

    return selected

# ==============================================================================
# 4) 섹션별 역할(Role Prompt) 및 작성 힌트(Queries)
# ==============================================================================

section_roles = {
    "연구개발 목표": "당신은 R&D PMO입니다. 단계/일괄 협약의 최종목표를 500자 내외로 명확·간결·정량화하여 작성합니다. 핵심 성능지표(KPI), 달성 기준(수치/단위/마일스톤), 검증 방법을 포함하고 모호한 표현은 배제합니다.",
    "연구개발 내용": "당신은 기술 총괄(Tech Lead)입니다. 전체 연구범위를 1,000자 내외로 구조화해 기술 요소, 서브태스크, 인터페이스, 데이터/시스템 흐름을 설명하고 표준·규격·평가계획을 명시합니다.",
    "연구개발성과 활용계획 및 기대효과": "당신은 사업전략/사업개발(BD) 담당자입니다. 수요처, 적용 시나리오, 도입·확산 경로, 수익/비용 구조, 경제적 파급효과를 500자 내외로 정량·정성 지표와 함께 제시합니다.",
    "연구기획과제의 개요": "당신은 제안서 총괄 에디터입니다. 목적·필요성·기대효과를 일관된 논리로 요약해 과제가 해결하는 문제와 중요성을 한눈에 보이게 작성합니다.",
    "연구개발과제의 배경": "당신은 정책/RFP 적합성 분석가입니다. 관련 선행연구·시장/기술 동향·정부 정책·RFP/품목요약서 부합성을 근거와 함께 정리하고 제안 맥락을 명확히 합니다.",
    "연구개발과제의 필요성": "당신은 산업분석가입니다. 현황·문제점·시장규모/성장률·규제 및 정책 요구를 데이터로 제시하고, 해결 필요성을 인과적으로 설득력 있게 제시합니다.",
    "보안등급의 분류 및 해당 사유": "당신은 보안관리 책임자입니다. 국가연구개발혁신법 시행령 제45조 및 산업기술혁신사업 보안관리요령 제9조 기준을 근거로 보안등급과 결정 사유를 간결히 기재합니다.",
    "기술개발 핵심어(키워드)": "당신은 표준/용어 관리자입니다. 과제의 핵심 용어 5개를 한글/영문 정식 명칭으로 제시하고, 표준(협회/학회) 정의에 부합하도록 작성합니다.",
    "연차별 개발목표": "당신은 일정/성과관리 PM입니다. 연차별(1년차~n년차) 목표를 기관(주관/공동/참여 연구원)별로 구분해 KPI·마일스톤·검증기준을 정량화하여 제시합니다.",
    "연차별 개발내용 및 범위": "당신은 공동연구 컨소시엄 코디네이터입니다. 기관별 역할·범위·인계·의존성을 명확히 기술하고 중복/누락 없이 연차별 산출물과 책임을 표로 정리합니다(공동기관 없으면 생략).",
    "추진방법 및 전략": "당신은 기술전략/실험 설계 책임자입니다. 방법론(데이터·알고리즘·장비), 리스크와 대응책, 실험/검증 계획(평가지표·샘플수·통계/검증 절차)을 구체적으로 기술합니다.",
    "과제 성과의 활용방안": "당신은 제품/사업화 매니저입니다. 성과의 적용 분야, 기술 파급효과, 에너지 절감·환경 개선 등 기술적·사회적 효익을 사용 시나리오와 함께 제시합니다.",
    "신규사업 신설의 기대효과": "당신은 전략기획 임원입니다. 시장 창출, 일자리, 수입대체, 수출 증대, 비용 절감 등 경제·산업적 효과를 정량 지표(금액, 비율, 기간)와 함께 제시합니다.",
    "사회적 가치 창출 계획": "당신은 ESG/사회가치 책임자입니다. 개요-비전-목표-세부계획-기대효과 체계로 13개 사회적 가치 범주와의 연계를 명확히 하고 측정 가능한 지표를 포함합니다.",
    "사회적 가치창출의 기대효과": "당신은 임팩트 평가자입니다. 보건·안전·포용·지역·환경·민주성 등 사회적 가치 지표를 중심으로 성과/파급효과를 정량·정성으로 제시합니다.",
    "경제적 성과창출의 기대효과": "당신은 재무 담당자입니다(기업 작성). 매출/원가/영업이익, ROI/NPV, 고용효과 등 재무적 성과 전망을 가정과 산식(간단) 포함하여 명료하게 제시합니다.",
    "신규 인력 채용 계획 및 활용 방안": "당신은 HR 책임자입니다. 신규/기존 채용 구분, 채용 시점·역할·배치·활용 계획, 역량 매핑과 교육/온보딩 계획을 일정표와 함께 제시합니다."
}

section_queries = {
    "연구개발 목표": "최종목표(단계/일괄 협약목표)를 과제의 연구기획목표를 500자 내외로 기재합니다.",
    "연구개발 내용": "전체내용을 1,000자 내외로 기재합니다.",
    "연구개발성과 활용계획 및 기대효과": "연구기획의 수요처, 활용내용, 경제적 파급효과 등을 500자 내외로 기재합니다(연구시설ㆍ장비 구축 과제일 경우 성과관리/자립운영/수입금 관리 계획 포함).",
    "연구기획과제의 개요": "연구개발과제의 목적, 필요성, 기대효과를 일관된 논리로 요약합니다.",
    "연구개발과제의 배경": "선행연구/시장·기술 동향/정책·RFP 부합성을 근거와 함께 기재합니다.",
    "연구개발과제의 필요성": "현황·문제점·시장규모/성장률·규제 및 정책 요구를 데이터로 제시합니다.",
    "보안등급의 분류 및 해당 사유": "법령/요령 기준을 근거로 보안등급과 결정 사유를 기재합니다.",
    "기술개발 핵심어(키워드)": "핵심 용어 5개(한글/영문) 정식 명칭으로 제시합니다.",
    "연차별 개발목표": "연차별 목표를 기관별로 KPI/마일스톤/검증기준과 함께 정량화합니다.",
    "연차별 개발내용 및 범위": "기관별 역할/범위/인계/의존성을 명확히 기술합니다. 명확한 기술의 예시는 가이드라인을 반드시 참고합니다.",
    "추진방법 및 전략": "방법론/리스크/대응책/실험·검증 계획을 구체화합니다.",
    "과제 성과의 활용방안": "성과의 적용 분야/파급효과/환경 개선 효과 등을 기술합니다.",
    "신규사업 신설의 기대효과": "시장/일자리/수출/비용절감 등 경제·산업적 효과를 정량 지표와 함께.",
    "사회적 가치 창출 계획": "개요-비전-목표-세부계획-기대효과 체계로 작성합니다.",
    "사회적 가치창출의 기대효과": "사회적 가치 지표 중심으로 정량·정성 성과를 제시합니다.",
    "경제적 성과창출의 기대효과": "기업만 작성. 매출/원가/ROI/고용효과 등을 제시합니다.",
    "신규 인력 채용 계획 및 활용 방안": "신규/기존 구분, 시점/배치/온보딩 계획을 포함합니다.",
}

# ==============================================================================
# 5) 생성 모델/토크나이저 로드 (생성용 토크나이저 사용 권장)
# ==============================================================================

# 생성 모델 전용 토크나이저 사용 
gen_tok = AutoTokenizer.from_pretrained(GEN_NAME, use_fast=True)
gen_tok.padding_side    = "left"   # causal LM에서 좌패딩이 유리한 경우가 많음
gen_tok.truncation_side = "left"

if gen_tok.pad_token_id is None:
    # PAD 토큰이 없으면 EOS를 패드로 재지정 (모델에 따라 필요)
    gen_tok.pad_token = gen_tok.eos_token

# 생성 LLM 로드 (CUDA면 fp16, 아니면 fp32)
gen_model = AutoModelForCausalLM.from_pretrained(
    GEN_NAME,
    trust_remote_code=True,
    torch_dtype=(torch.float16 if DEVICE == "cuda" else torch.float32),
    device_map="auto"
).eval()

if hasattr(gen_model, "generation_config"):
    gen_model.generation_config.pad_token_id = gen_tok.pad_token_id

# 간편 파이프라인 (문자열 리스트도 입력 가능)
gen_pipe = pipeline("text-generation", model=gen_model, tokenizer=gen_tok)

# ==============================================================================
# 6) 컨텍스트/프롬프트 구성
# ==============================================================================

def build_ctx_block(section: str, project_name: str, keywords: str, topk: int = TOPK) -> str:
    """
    섹션/프로젝트/키워드 → E5 쿼리 임베딩 → FAISS 검색(POOL=N) → MMR 다양성 리랭킹 → 상위 K 스니펫 컨텍스트 구성.
    """
    # 1) 쿼리 텍스트 구성(섹션 역할/요구를 한 줄로 요약 + 프로젝트/키워드 결합)
    role_hint  = section_roles.get(section, "")
    query_hint = section_queries.get(section, "")
    base_query = ' | '.join(filter(None, [section, project_name, keywords, role_hint, query_hint]))
    base_query = " ".join(base_query.split())[:1024]  # 과도한 길이 방어

    # 2) 인덱스 없음 → fallback: combined_text 앞부분
    if faiss_index is None or faiss_index.ntotal == 0:
        return f"- [RAG 근거] {first_n_chars(combined_text)}" if combined_text.strip() else "- [근거] 자료 없음"

    # 3) 1차 검색(POOL) → MMR로 K개 선별
    q_emb = embed_texts_e5([base_query], is_query=True)  # (1, d)
    pool_n = min(SEARCH_POOL, faiss_index.ntotal)
    D, I = faiss_index.search(q_emb, pool_n)             # (1, pool_n)

    cand_idx = I[0]
    cand_idx = cand_idx[cand_idx >= 0]
    if cand_idx.size == 0:
        return f"- [RAG 근거] {first_n_chars(combined_text)}" if combined_text.strip() else "- [근거] 자료 없음"

    cand_emb = doc_emb[cand_idx]  # (pool_n, d)
    sel_rel = mmr_select(q=q_emb[0], docs=cand_emb, k=topk, lam=MMR_LAMBDA)
    final_idx = [int(cand_idx[i]) for i in sel_rel]

    # 4) 컨텍스트 라인 구성
    lines = []
    for rank, idx in enumerate(final_idx, start=1):
        src = doc_texts[idx]
        lines.append(f"- [근거 {rank}] {first_n_chars(src)}")
    return "\n".join(lines) if lines else "- [근거] 자료 미발견"

def build_prompt(section: str, project_name: str, depart_name: str, project_no: str,
                 period: str, budget: str, ctx_block: str) -> str:
    """
    섹션별 역할(Role) + 작성조건 + 근거텍스트를 하나의 시스템 프롬프트로 구성.
    """
    role_instruction = section_roles.get(section, "당신은 연구자이자 계획서 및 해당 분야 전문가입니다.")
    request_hint     = section_queries.get(section, "문장은 ~함. ~음. 명사 등 으로 마무리하고, 각 문단마다 핵심 키워드 포함, 문장 길이/시작 다양화, 중복 표현은 피한다.")

    prompt = f"""
# #=========== 자동 문장 생성
# 역할: {role_instruction}
# 작성 항목: [{section}]
# 세부사업명: {depart_name}
# 연구개발 과제번호: {project_no}
# 연구개발과제명: {project_name}
# 전체 연구개발기간: {period}
# 예산: {budget} 천원

# 작성 조건:
#     - 제시된 {GUIDELINE_FILE} 가이드라인을 엄격히 준수하여 작성함.
#     - 문장은 ~함, ~음, 명사로 마무리함(문장 종결 통일).
#     - 기술적 연관성이 낮은 미사여구 배제.
#     - 구체적인 규격/범위 포함.
#     - 아래 근거를 반영함:
# {ctx_block}
#     - 반드시 {RAG_JSON_FILES}의 작성 방식과 구성을 참고하여 작성함.
#     - 문단마다 핵심 키워드 포함, 문장 길이/시작 다양화, 중복 표현 회피.
#     - 전문적이면서 친화적인 톤.
#     - 모든 전문용어/약어에는 주석(full name) 표기.
# 요청된 형식(힌트): {request_hint}

#=========== 출력
""".strip()
    return prompt

# ==============================================================================
# 7) 생성 루틴
# ==============================================================================

def generate_sections_pipeline_batched(prompts: List[str], batch_size: int = 2) -> List[str]:
    """
    HF pipeline으로 문자열 리스트를 배치 생성.
    - 메모리 안전을 위해 기본 batch_size=2 (OOM 시 1로 낮추기)
    """
    outputs = []
    bat = batch_size
    for i in tqdm(range(0, len(prompts), bat), desc="Generate (pipeline/batch)"):
        chunk = prompts[i:i+bat]
        # pipe는 디바이스 자동 관리, to(device) 필요 없음
        out = gen_pipe(
            chunk,
            batch_size=bat,
            max_new_tokens=GEN_MAX_NEW_TOKENS,
            do_sample=GEN_DO_SAMPLE,
            temperature=(GEN_TEMPERATURE if GEN_DO_SAMPLE else None)
        )
        # pipeline 반환 형태(list of list of dict) 정규화
        for item in out:
            # item은 [ {'generated_text': ...} ] 형태이므로 item[0] 접근
            outputs.append(item[0]["generated_text"]) 
    return outputs

def collate_pad(batch, pad_token_id: int) -> Dict[str, torch.Tensor]:
    """
    DataLoader용 수동 collate: 배치 내 최대 길이에 맞춰 좌패딩/어텐션마스크 생성.
    """
    max_len = max(len(x["input_ids"]) for x in batch)
    input_ids, attn = [], []
    for x in batch:
        ids = x["input_ids"]
        am  = x["attention_mask"]
        pad_len = max_len - len(ids)
        input_ids.append(ids + [pad_token_id] * pad_len)
        attn.append(am + [0] * pad_len)
    return {
        "input_ids": torch.tensor(input_ids, dtype=torch.long),
        "attention_mask": torch.tensor(attn, dtype=torch.long)
    }

def generate_sections_datasets_manual(prompts: List[str], batch_size: int = 2) -> List[str]:
    """
    datasets + DataLoader + model.generate 경로 (세밀 제어).
    - 기본 batch_size=2 (OOM 시 1로)
    - num_workers=0 (환경 호환성↑, 메모리 안전)
    """
    ds = Dataset.from_dict({"prompt": prompts})

    def tok_fn(batch):
        enc = gen_tok(
            batch["prompt"],
            padding=False,           # collate에서 패딩
            truncation=True,         # 프롬프트 과도 길이 방지
            return_attention_mask=True
        )
        return enc

    ds_tok = ds.map(tok_fn, batched=True, batch_size=1024, remove_columns=["prompt"])
    ds_tok.set_format(type=None)

    pad_id = gen_tok.pad_token_id
    # pad_token 설정은 로드 시 이미 했지만, 안전을 위해 다시 확인
    if pad_id is None:
        gen_tok.pad_token = gen_tok.eos_token
        pad_id = gen_tok.pad_token_id

    loader = DataLoader(
        ds_tok,
        batch_size=batch_size,
        shuffle=False,
        num_workers=0,           # 안전 기본값
        pin_memory=True,         # OOM 시 False 고려
        persistent_workers=False,
        collate_fn=lambda b: collate_pad(b, pad_id)
    )

    model_device = next(gen_model.parameters()).device
    gen_out_texts: List[str] = []

    torch.set_grad_enabled(False)
    for batch in tqdm(loader, desc="Generate (datasets/manual)"):
        input_ids      = batch["input_ids"].to(model_device, non_blocking=True)
        attention_mask = batch["attention_mask"].to(model_device, non_blocking=True)

        gen_kwargs = dict(
            max_new_tokens=GEN_MAX_NEW_TOKENS,
            do_sample=GEN_DO_SAMPLE
        )
        if GEN_DO_SAMPLE and GEN_TEMPERATURE is not None:
            gen_kwargs["temperature"] = GEN_TEMPERATURE

        outputs = gen_model.generate(
            input_ids=input_ids,
            attention_mask=attention_mask,
            **gen_kwargs
        )
        texts = gen_tok.batch_decode(outputs, skip_special_tokens=True)
        gen_out_texts.extend(texts)

    return gen_out_texts

def generate_all_sections_batched(sections: List[str], project_name: str, depart_name: str,
                                  project_no: str, period: str, budget: str, keywords: str) -> Dict[str, str]:
    """
    섹션 순회 → 컨텍스트 검색/구성 → 프롬프트 생성 → 배치 생성 → 섹션→텍스트 매핑 반환.
    """
    prompts = []
    for section in sections:
        ctx_block = build_ctx_block(section, project_name, keywords, topk=TOPK)
        prompt    = build_prompt(section, project_name, depart_name, project_no, period, budget, ctx_block)
        prompts.append(prompt)

    if GEN_MODE == "pipeline_batch":
        texts = generate_sections_pipeline_batched(prompts, batch_size=2)
    elif GEN_MODE == "datasets_manual":
        texts = generate_sections_datasets_manual(prompts, batch_size=2)
    else:
        raise ValueError(f"Unknown GEN_MODE: {GEN_MODE}")

    return {sec: txt for sec, txt in zip(sections, texts)}

# ==============================================================================
# 8) DOCX 렌더링
# ==============================================================================

def render_doc(project_name: str, depart_name: str, project_no: str, period: str, budget: str, keywords: str) -> None:
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')  # 한글 폰트 매핑
    font.size = Pt(11)

    # 메타 정보 헤더
    doc.add_heading("연구개발계획서", 0)
    doc.add_paragraph(f"세부사업명: {depart_name}")
    doc.add_paragraph(f"연구개발 과제번호: {project_no}")
    doc.add_paragraph(f"연구개발과제명: {project_name}")
    doc.add_paragraph(f"전체 연구개발기간: {period}")
    doc.add_paragraph(f"예산: {budget} 천원")
    doc.add_paragraph("")

    sections = list(section_roles.keys())

    # 섹션별 텍스트 생성 (검색 RAG 사용)
    section_texts = generate_all_sections_batched(
        sections, project_name, depart_name, project_no, period, budget, keywords
    )

    # 섹션별 Heading + 본문 추가
    for section in sections:
        doc.add_heading(section, level=1)
        raw_text    = section_texts[section]              # 전체(프롬프트+출력)가 포함될 수 있음
        cleaned_txt = clean_generated_text(raw_text)  # "#=========== 출력" 이후만 취하고 마크다운 제거
        doc.add_paragraph(cleaned_txt)

    outpath = "e5base_계획서.docx"
    doc.save(outpath)
    print(f"[DONE] {outpath} 생성 완료 (GEN_MODE={GEN_MODE})")

# ==============================================================================
# 9) 실행 예시
# ==============================================================================

if __name__ == "__main__":
    project_name = "VUNO Med-Chest X-ray"
    depart_name  = "산업기술R&D연구기획사업"
    project_no   = "123456789"
    period       = "2023. 6. 1 ~"
    budget       = "-"
    keywords     = "흉부 엑스레이(Chest X-ray), 인공지능 의료영상(AI Medical Imaging), 폐질환 진단(Lung Disease Diagnosis), 의료 인공지능 소프트웨어(Medical AI Software), 산업기술 R&D (Industrial Technology R&D)"

    # 메모리 여유 없으면 아래 값 조절:
    # - GEN_MAX_NEW_TOKENS = 512
    # - generate_* batch_size = 1
    # - DataLoader pin_memory=False
    try:
        render_doc(project_name, depart_name, project_no, period, budget, keywords)
    except Exception as e:
        print(f"\n 파이프라인 실행 중 오류 발생. 외부 파일 경로({GUIDELINE_FILE}, {RAG_JSON_FILES})나 환경(GPU/메모리)을 확인하세요.")
        print(f"오류 상세: {e}")
