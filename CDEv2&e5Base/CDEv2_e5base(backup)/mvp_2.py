from docx import Document
from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline
from docx.shared import Pt
from docx.oxml.ns import qn

# ===== [추가] RAG 유틸 의존성 =====
import fitz  # PyMuPDF
import re
from sentence_transformers import SentenceTransformer
import faiss
from rank_bm25 import BM25Okapi

# ===== [추가] PDF → 텍스트 정제 =====
def load_pdf_text(path: str) -> str:
    doc = fitz.open(path)
    pages = []
    for p in doc:
        text = p.get_text("text")
        text = re.sub(r"\s{2,}", " ", text)
        pages.append(text)
    return "\f".join(pages)

def clean_text(full_text: str) -> str:
    pages = full_text.split("\f")
    cleaned = []
    for pg in pages:
        pg = re.sub(r"전자공시시스템.*?(Page\s*\d+)?", "", pg, flags=re.IGNORECASE)
        pg = re.sub(r"\s{2,}", " ", pg)
        cleaned.append(pg.strip())
    return "\n".join(cleaned)

# ===== [추가] 섹션 앵커 분할 =====
ANCHOR_PATTERNS = [
    r"^\s*[IVXLC]+\.\s.+$",
    r"^\s*\d+\.\s.+$",
    r"^\s*\d+\-\d+\.\s.+$",
    r"^\s*제\s*\d+\s*기.*$",
    r"^\s*\(\d+\)\s.+$",
]

def detect_anchors(text: str):
    lines = text.splitlines()
    anchors = []
    for i, line in enumerate(lines):
        for pat in ANCHOR_PATTERNS:
            if re.match(pat, line.strip()):
                anchors.append((i, line.strip()))
                break
    return anchors

def segment_by_anchors(text: str):
    lines = text.splitlines()
    anchors = detect_anchors(text)
    if not anchors:
        return [{"title": "FULL", "text": text}]
    segments = []
    for idx, (lineno, title) in enumerate(anchors):
        start = lineno
        end = anchors[idx+1][0] if idx+1 < len(anchors) else len(lines)
        seg_text = "\n".join(lines[start:end]).strip()
        segments.append({"title": title, "text": seg_text})
    return segments

# ===== [추가] 인덱서/검색기 =====
class TextIndexer:
    def __init__(self, model_name="intfloat/multilingual-e5-base"):
        self.model = SentenceTransformer(model_name)
        self.idx = None
        self.chunks = []

    def chunk(self, segments, max_chars=1200):
        res = []
        for seg in segments:
            buf, acc = [], 0
            for line in seg["text"].split("\n"):
                if not line.strip():
                    continue
                if acc + len(line) > max_chars and buf:
                    res.append({"title": seg["title"], "text": " ".join(buf)})
                    buf, acc = [], 0
                buf.append(line.strip()); acc += len(line)
            if buf:
                res.append({"title": seg["title"], "text": " ".join(buf)})
        self.chunks = res
        return res

    def build(self):
        texts = [c["text"] for c in self.chunks]
        if not texts:
            self.idx = None
            return
        embs = self.model.encode(texts, normalize_embeddings=True)
        dim = embs.shape[1]
        self.idx = faiss.IndexFlatIP(dim)
        self.idx.add(embs.astype("float32"))

    def search_dense(self, query: str, topk=4):
        if self.idx is None or not self.chunks:
            return []
        qv = self.model.encode([query], normalize_embeddings=True)
        D, I = self.idx.search(qv, topk)
        items = []
        for i in I[0]:
            if i < 0:
                continue
            items.append(self.chunks[i])
        return items

class HybridRetriever:
    def __init__(self, dense_indexer):
        self.dense = dense_indexer
        self.corpus = [c["text"] for c in self.dense.chunks]
        self.bm25 = BM25Okapi([t.split() for t in self.corpus]) if self.corpus else None

    def search(self, query, k_dense=4, k_bm25=4):
        dense_hits = self.dense.search_dense(query, topk=k_dense)
        bm_hits = []
        if self.bm25 is not None:
            bm_ids = self.bm25.get_top_n(query.split(), list(range(len(self.corpus))), n=k_bm25)
            bm_hits = [self.dense.chunks[i] for i in bm_ids]
        seen = set()
        results = []
        for it in dense_hits + bm_hits:
            key = it["text"][:80]
            if key in seen: 
                continue
            seen.add(key)
            results.append(it)
        return results[:max(k_dense, k_bm25)]

# ===============================
# 1. 모델 로드 (기존 유지, 옵션만 살짝 조정)
# ===============================
model_name = "skt/A.X-4.0-Light"
tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModelForCausalLM.from_pretrained(
    model_name,
    device_map="auto",
    torch_dtype="auto"
)
generator = pipeline("text-generation", model=model, tokenizer=tokenizer)

# ===============================
# 1-1. [추가] PDF 경로 입력 및 RAG 인덱스 준비
# ===============================
pdf_path = input("사업보고서 PDF 경로: ").strip()
raw_text = load_pdf_text(pdf_path)
cleaned_text = clean_text(raw_text)
segments = segment_by_anchors(cleaned_text)
indexer = TextIndexer()
chunks = indexer.chunk(segments, max_chars=1200)
indexer.build()
retriever = HybridRetriever(indexer)

# ===============================
# 2. 사용자 입력 받기 (기존)
# ===============================
project_name = input("사업명: ")
company_name = input("회사명: ")
manager_name = input("담당자: ")
keywords = input("문제 정의 키워드 (쉼표로 구분): ")
budget = input("예산(단위: 백만원): ")

# ===============================
# 3. 섹션별 롤플레이 프롬프트 (기존 + 약간 보강)
# ===============================
section_roles = {
    "사업 개요": "당신은 투자 심사역입니다. 투자자의 관점에서 사업 개요를 설명하세요.",
    "문제 정의": "당신은 현업 부서장입니다. 실제 업무에서 느끼는 문제 정의를 강조하세요.",
    "해결 방안": "당신은 해당 분야의 전문적인 분석가입니다. 기술적 해결 방안을 구체적으로 제시하세요.",
    "기대 효과": "당신은 CEO입니다. 경영적 기대 효과를 전략적 가치 중심으로 설명하세요."
}

# ===============================
# 3-1. [추가] 섹션별 RAG 질의 템플릿
# ===============================
section_queries = {
    "사업 개요": "회사 개요 설립일 본점 소재지 주된 사업 요약",
    "문제 정의": f"{keywords} 관련 당사 사업환경 리스크 또는 문제 서술",
    "해결 방안": f"{keywords} 해결 기술/제품/서비스 방향 핵심 근거",
    "기대 효과": "경영성과, 시장성, 경쟁우위, 수익성, 전략적 기대효과"
}

# ===============================
# 4. 자동 문장 생성 함수 (RAG 근거 포함으로 변경)
# ===============================
def search_contexts(section: str, topk=4):
    query = section_queries.get(section, section)
    hits = retriever.search(query, k_dense=4, k_bm25=4)
    # 너무 길면 잘라내기
    contexts = []
    max_ctx_len = 900  # 토큰 보호용 대략적 글자 제한
    for h in hits:
        t = h["text"]
        if len(t) > max_ctx_len:
            t = t[:max_ctx_len] + "..."
        contexts.append(t)
    return contexts

def build_prompt_with_context(section, role_instruction, contexts):
    ctx_block = "\n\n".join([f"[근거]\n{c}" for c in contexts]) if contexts else "[근거]\n(해당 섹션에 대한 근거 스니펫 없음)"
    prompt = f"""
역할: {role_instruction}
작성 항목: [{section}]
회사명: {company_name}
사업명: {project_name}

작성 조건:
- 아래 근거 텍스트 안에서만 답변할 것
- 출처 밖 가정 금지, 누락 정보는 '확인 필요'로 표기
- 수치/날짜/명칭은 원문 그대로 유지
- 국가사업 제안서 톤
- 3~5문장
- 다른 항목과 중복 표현 최소화

{ctx_block}
"""
    # 파이프라인용 프롬프트는 한 번에 전달
    return prompt.strip()

def generate_text(section, keywords=""):
    role_instruction = section_roles.get(section, "")
    contexts = search_contexts(section, topk=4)
    prompt = build_prompt_with_context(section, role_instruction, contexts)
    output = generator(
        prompt,
        max_new_tokens=220,
        do_sample=False,      # 사실성 위해 샘플링 끔
        temperature=0.2,      # 낮은 온도로 보수적 생성
        top_p=0.9,
        repetition_penalty=1.05,
        eos_token_id=tokenizer.eos_token_id
    )
    # pipeline은 프롬프트까지 포함해서 반환할 수 있으므로 후처리
    text = output[0]["generated_text"]
    # 생성부만 추정 추출 (간단히 프롬프트 길이 기준)
    gen = text[len(prompt):].strip() if text.startswith(prompt) else text.strip()
    return gen

# ===============================
# 5. DOCX 문서 생성 (기존)
# ===============================
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
font.size = Pt(11)

doc.add_heading("국가사업 제안서 (간략 버전)", 0)
doc.add_paragraph(f"사업명: {project_name}")
doc.add_paragraph(f"회사명: {company_name}")
doc.add_paragraph(f"담당자: {manager_name}")
doc.add_paragraph(f"예산: {budget} 백만원")
doc.add_paragraph("")

sections = ["사업 개요", "문제 정의", "해결 방안", "기대 효과"]

for section in sections:
    doc.add_heading(section, level=1)
    doc.add_paragraph(generate_text(section, keywords))

doc.add_page_break()

# ===============================
# 6. 파일 저장 (기존)
# ===============================
output_file = "mvp_proposal_ai.docx"
doc.save(output_file)
print(f"✅ '{output_file}' 파일이 생성되었습니다!")