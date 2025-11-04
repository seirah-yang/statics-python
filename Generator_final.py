# ================================================================
# RND_Plan_Pipeline.py
# SKT A.X 4.0 + E5-Large + XLM-RoBERTa NLI 기반 문서 생성 파이프라인
# ================================================================

# -*- coding: utf-8 -*-
"""
(A) ContextSearch
        └─ e5-large 기반 근거 검색 (가중치 반영)
(B) DraftGeneration
        └─ A.X 4.0 Light으로 초안 생성
(C) Validation
        ├─ Rule-based 검사
        └─ NLI entailment 점수 계산
(D) Export Report
        └─ DOCX 저장 + 검증리포트 + 근거문서목록
"""

# ================================================================
# 1. 기본 라이브러리 로드
# ================================================================
import os, re, json, glob, torch
import numpy as np
from tqdm.auto import tqdm
from transformers import (
    AutoTokenizer, AutoModelForCausalLM, pipeline,
    AutoModelForSequenceClassification
)
from sentence_transformers import SentenceTransformer, util
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from PyPDF2 import PdfReader


# ================================================================
# 2. 환경 변수 설정
# ================================================================
LAW_DIR = "./reference_file"  # 참고문헌 디렉토리
DEVICE = "cuda" if torch.cuda.is_available() else "cpu"

E5_NAME = "intfloat/multilingual-e5-large"
GEN_NAME = "skt/A.X-4.0-Light"
NLI_MODEL = "joeddav/xlm-roberta-large-xnli"

GEN_MAX_NEW_TOKENS = 1500
GEN_DO_SAMPLE = False

torch.cuda.empty_cache()

# ================================================================
# 3. 헬퍼 함수
# ================================================================
def first_n_lines(text: str, n_chars=400):
    t = " ".join(str(text).split())
    return t[:n_chars]

def clean_generated_text(text: str) -> str:
    text = re.sub(r"[•●▪▶◇◆□▪️▫️–]", " ", text)
    text = re.sub(r"^\s*[-#*]+\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"\s+([\.,;:])", r"\1", text)
    return text.strip()

def has_number(text: str):
    return bool(re.search(r"\d", text))

# ================================================================
# 4. 섹션 정의 (Full Restoration)
# ================================================================
sections = [
    {"section": "연구기획과제의 개요",
        "role": "제안서 총괄 에디터",
        "query": "과제의 목적, 필요성, 기대효과를 논리적으로 연결하여 핵심 내용을 500자 내외로 명료하게 요약함. 평가자가 과제의 전체 구조를 한눈에 이해할 수 있도록 기술함.",
        "constraints": ["500자 내외", "논리적 연결 구조 유지", "핵심 요약 중심"],
        "output_format": "서술문"},

    {"section": "연구개발과제의 배경",
        "role": "R&D 기획 전문가",
        "query": "제공된 선행연구, 시장 동향, 정책자료를 근거로 본 과제가 추진되어야 하는 당위성과 RFP(공고문) 부합성을 논리적으로 제시함.",
        "constraints": ["데이터 근거 포함", "RFP 문항 부합성 명시"],
        "output_format": "서술문"},

    {"section": "연구개발과제의 필요성",
        "role": "산업분석가",
        "query": "데이터와 사례를 근거로 현재 기술적·산업적 문제점을 제시하고, 본 과제가 이를 해결해야 하는 필요성을 인과적으로 서술함.",
        "constraints": ["인과관계 구조", "데이터 근거 제시"],
        "output_format": "서술문"},

    {"section": "기술개발 핵심어(키워드)",
        "role": "기술 네이밍 전략가",
        "query": "과제의 정체성과 핵심 기술을 대표하는 키워드 5개를 국문·영문 공식 명칭으로 제시함. 각 용어는 국제 표준 또는 학술 정의를 근거로 하며, 선정 사유를 간단히 명시함.",
        "constraints": ["국문·영문 병기", "국제 표준 기반 정의 포함", "5개 이내"],
        "output_format": "표(키워드/영문명/정의/출처)"},

    {"section": "연구개발 목표",
        "role": "R&D PMO",
        "query": "과제의 최종 목표를 500자 내외로 명확히 기술함. 핵심 성능지표(KPI), 달성 기준(수치·단위·마일스톤), 검증 방법을 포함하며, 모호한 표현은 사용하지 않음.",
        "constraints": ["정량화된 수치 포함", "KPI 및 검증방법 명시"],
        "output_format": "서술문 + 표(KPI/단위/기준/검증방법)"},

    {"section": "연구개발 내용",
        "role": "기술총괄자(Tech Lead)",
        "query": "전체 연구 범위를 1,000자 내외로 체계적으로 기술함. 핵심 기술요소, 세부 과제 구조, 데이터 및 시스템 흐름, 성능 평가 계획을 명시함.",
        "constraints": ["1,000자 내외", "기술요소 및 평가계획 포함"],
        "output_format": "서술문 + 도식(기술흐름)"},

    {"section": "연차별 개발목표",
        "role": "PMO 리더",
        "query": "최종 목표 달성을 위한 연차별 및 기관별 개발 목표를 정량적으로 제시함. 각 연차별 KPI, 마일스톤, 검증 방법을 명확히 포함함.",
        "constraints": ["연차별 구분", "정량적 지표 포함"],
        "output_format": "표(연차/KPI/마일스톤/검증방법)"},

    {"section": "연차별 개발내용 및 범위",
        "role": "공동연구 컨소시엄 코디네이터",
        "query": "참여 기관별 역할, 책임, 연차별 산출물을 중복 및 누락 없이 기술함. 공동기관이 없는 경우 해당 항목은 생략함.",
        "constraints": ["기관별 역할 명확화", "중복/누락 금지"],
        "output_format": "표(기관/역할/산출물/책임)"},

    {"section": "추진방법 및 전략",
        "role": "총괄 아키텍트(Chief Architect)",
        "query": "핵심 기술개발 방법론, 예측 가능한 리스크와 대응 방안, 성능 검증 계획을 논리적으로 제시함. 기술의 우수성과 실현 가능성을 입증함.",
        "constraints": ["핵심 방법론 포함", "리스크 및 검증계획 명시"],
        "output_format": "서술문 + 표(리스크/대응방안)"},

    {"section": "과제 성과의 활용방안",
        "role": "사업개발 총괄(BD Head)",
        "query": "연구 성과가 실제 산업 및 시장에서 어떻게 활용될 수 있는지를 구체적으로 제시함. 목표 시장, 주요 수요처, 핵심 적용 시나리오, 기술의 차별화된 가치를 중심으로 사업화 방향을 설명함.",
        "constraints": ["시장 시나리오 포함", "Value Proposition 중심"],
        "output_format": "서술문 + 표(시장/수요처/적용시나리오)"},

    {"section": "신규사업 신설의 기대효과",
        "role": "거시경제 분석가(Macro-Economic Analyst)",
        "query": "본 과제가 국가 경제에 미치는 파급효과를 정량적 지표로 제시함. 시장 창출, 수입 대체, 수출 증대, 일자리 창출 등 거시적 효과를 수치로 증명함.",
        "constraints": ["정량적 수치 기반", "경제효과 명시"],
        "output_format": "표(지표/예상값/근거자료)"},

    {"section": "사회적 가치 창출 계획",
        "role": "사회적 가치 전략가(Social Value Strategist)",
        "query": "과제의 사회적 비전과 목표를 정의하고, 이를 달성하기 위한 구체적인 실행 로드맵을 수립함. 보건, 환경, 안전 등 사회적 가치 범주와 연계함.",
        "constraints": ["사회적 가치 범주 명시", "로드맵 포함"],
        "output_format": "서술문 + 표(목표/실행단계/성과지표)"},

    {"section": "사회적 가치창출의 기대효과",
        "role": "임팩트 평가 전문가(Impact Assessor)",
        "query": "사회적 가치 창출 계획이 실행되었을 때 예상되는 긍정적 변화를 정량적 및 정성적 임팩트 지표로 제시함. 사회적 파급효과를 객관적으로 설명함.",
        "constraints": ["정량/정성 지표 병기", "사회적 파급효과 포함"],
        "output_format": "표(지표유형/성과/측정방법)"},

    {"section": "경제적 성과창출의 기대효과",
        "role": "최고재무책임자(CFO)",
        "query": "기업 관점에서 본 과제의 재무적 성과를 구체적 수치와 함께 제시함. 예상 매출, 이익, 투자수익률(ROI), 순현재가치(NPV) 등 주요 지표를 근거와 함께 명료하게 기술함.",
        "constraints": ["재무지표 포함", "산출근거 명시"],
        "output_format": "표(지표/예상값/근거)"},

    {"section": "신규 인력 채용 계획 및 활용 방안",
        "role": "전략적 인사 파트너(Strategic HR Partner)",
        "query": "과제 수행에 필요한 핵심 인력의 채용, 배치, 교육 계획을 타임라인과 함께 제시함. 인력 확보 및 역량 극대화 방안을 구체적으로 기술함.",
        "constraints": ["타임라인 포함", "역량 강화 계획 명시"],
        "output_format": "표(직무/채용시점/교육계획)"},

    {"section": "보안등급의 분류 및 해당 사유",
        "role": "보안관리 책임자(Security Manager)",
        "query": "관련 법령 및 보안관리요령을 근거로 본 과제의 보안등급을 분류하고, 그 결정 사유를 간결하고 명확하게 기술함.",
        "constraints": ["법령 근거 포함", "사유 명시"],
        "output_format": "서술문"}
]

# ================================================================
# 5. 임베딩 기반 RAG 인덱스 구축
# ================================================================
CHUNK_MAX = 500
CHUNK_OVERLAP = 50
TOPK = 5

REF_WEIGHTS = {
    "행정업무의 운영 및 혁신에 관한 규정": 0.2,
    "국가연구개발사업 연구개발계획서": 0.2,
    "산업기술 R&D 과제명 작성 가이드라인": 0.2,
    "전략계획서 작성안내서": 0.1,
    "Vertical": 0.1
}
DEFAULT_WEIGHT = 0.05

def guess_weight(filename: str) -> float:
    for k, w in REF_WEIGHTS.items():
        if k.lower() in filename.lower():
            return w
    return DEFAULT_WEIGHT

def chunk_text(txt: str, max_chars=CHUNK_MAX, overlap=CHUNK_OVERLAP):
    txt = " ".join(str(txt).split())
    chunks = []
    i = 0
    while i < len(txt):
        j = min(len(txt), i + max_chars)
        chunks.append(txt[i:j])
        if j == len(txt): break
        i = max(0, j - overlap)
    return chunks

def load_reference_chunks(law_dir: str):
    items = []
    for p in glob.glob(os.path.join(law_dir, "*")):
        text = ""
        try:
            ext = p.split(".")[-1].lower()
            if ext == "pdf":
                reader = PdfReader(p)
                for pg in reader.pages:
                    text += (pg.extract_text() or "") + "\n"
            elif ext in ("docx", "doc"):
                d = Document(p)
                text = "\n".join([x.text for x in d.paragraphs])
            elif ext in ("rtf", "txt"):
                with open(p, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
        except Exception:
            continue
        base = os.path.basename(p)
        weight = guess_weight(base)
        for idx, ck in enumerate(chunk_text(text)):
            items.append({"source": base, "chunk_id": idx, "weight": weight, "text": ck})
    return items

print("[INFO] Building embedding index...")
embed_model = SentenceTransformer(E5_NAME, device=DEVICE)
REF_ITEMS = load_reference_chunks(LAW_DIR)
REF_EMBS = embed_model.encode(
    [it["text"] for it in REF_ITEMS],
    convert_to_tensor=True,
    normalize_embeddings=True
)
print(f"[INFO] Loaded {len(REF_ITEMS)} reference chunks.")


def search_reference(query: str, topk: int = TOPK):
    q = embed_model.encode(query, convert_to_tensor=True, normalize_embeddings=True)
    sims = util.cos_sim(q, REF_EMBS)[0]
    scores = [(float(s) * (1.0 + REF_ITEMS[i]["weight"]), i) for i, s in enumerate(sims)]
    scores.sort(key=lambda x: x[0], reverse=True)
    picks = []
    for sc, idx in scores[:topk]:
        it = REF_ITEMS[idx]
        picks.append({
            "source": it["source"],
            "chunk_id": it["chunk_id"],
            "weight": it["weight"],
            "score": round(sc, 4),
            "snippet": first_n_lines(it["text"], 400)
        })
    return picks


def format_ctx_block(refs):
    return "\n".join([
        f"- [{r['source']} | w={r['weight']:.2f} | score={r['score']:.3f}]\n  {r['snippet']}"
        for r in refs
    ])


# ================================================================
# 6. 모델 로드
# ================================================================
print("[INFO] Loading generation & validation models...")
gen_tok = AutoTokenizer.from_pretrained(GEN_NAME, use_fast=False, trust_remote_code=True)
if gen_tok.pad_token_id is None:
    gen_tok.pad_token = gen_tok.eos_token

gen_model = AutoModelForCausalLM.from_pretrained(GEN_NAME, trust_remote_code=True, device_map="auto").eval()
gen_pipe = pipeline("text-generation", model=gen_model, tokenizer=gen_tok)

nli_tok = AutoTokenizer.from_pretrained(NLI_MODEL)
nli_model = AutoModelForSequenceClassification.from_pretrained(NLI_MODEL).to(DEVICE).eval()


def nli_entail_vs_contra(premise, hypothesis):
    inputs = nli_tok(premise, hypothesis, return_tensors="pt", truncation=True, padding=True).to(DEVICE)
    with torch.no_grad():
        logits = nli_model(**inputs).logits
    probs = torch.softmax(logits, dim=-1).squeeze(0).tolist()
    entail, neutral, contra = probs[0], probs[1], probs[2]
    return entail, contra


# ================================================================
# 7. 검증 및 생성
# ================================================================
def validate_output(section_obj, generated_text: str, refs):
    report = []
    cleaned = generated_text.strip()
    maxlen = 10**9
    for c in section_obj.get("constraints", []):
        m = re.search(r"(\d+)\s*자", c)
        if m:
            maxlen = int(m.group(1)) + int(0.3*int(m.group(1)))
            break
    if len(cleaned) > maxlen:
        report.append(f" 길이 초과: {len(cleaned)}자 > 허용 {maxlen}자")
    if any(k in section_obj["section"] for k in ["목표", "KPI"]) or "정량" in " ".join(section_obj.get("constraints", [])):
        if not has_number(cleaned):
            report.append("정량 지표(숫자) 미포함")
    entail_sum, contra_sum = 0.0, 0.0
    for r in refs[:3]:
        e, c = nli_entail_vs_contra(r["snippet"], cleaned)
        entail_sum += e; contra_sum += c
    score = entail_sum / (entail_sum + contra_sum + 1e-6)
    report.append(f" NLI 정합도: {score:.2f}")
    if not report:
        report.append(" PASS-validation")
    return report
  
# 제 1원칙 = 역할: section_obj['role'] & 제 2원칙 = 아래 근거 기반으로 작성 {ctx_block}
# 제3원칙(출력 구조 고정화) 선택이지만 권장 - 문서의 형식을 일정하게 유지 
def build_prompt(section_obj, project_name, depart_name, project_no, period, budget):
    refs = search_reference(section_obj["query"], topk=TOPK)
    ctx_block = format_ctx_block(refs)
    return f"""
# 역할: {section_obj['role']}
# 작성 항목: [{section_obj['section']}]
# 세부사업명: {depart_name}
# 과제번호: {project_no}
# 과제명: {project_name}
# 기간: {period}
# 예산: {budget} 천원

# 작성 조건:
# - 문체: '~함, ~음, ~됨' 보고서형 서술체
# - {', '.join(section_obj['constraints'])} 준수
# - 아래 근거문서 기반으로 작성
{ctx_block}

# 요청 내용:
# {section_obj['query']}

#출력형식 : 
# {section_obj['output_format']}
#=========== 출력
""".strip()

def generate_text_batch(prompts, batch_size=1):
    outputs = []
    for i in tqdm(range(0, len(prompts), batch_size), desc="Generate"):
        chunk = prompts[i:i+batch_size]
        out = gen_pipe(chunk, batch_size=batch_size, max_new_tokens=GEN_MAX_NEW_TOKENS, do_sample=GEN_DO_SAMPLE)
        for o in out:
            txt = o[0]["generated_text"] if isinstance(o, list) else o["generated_text"]
            outputs.append(txt)
    return outputs


# ================================================================
# 8. DOCX 생성
# ================================================================
def add_reference_section(doc, all_refs):
    doc.add_page_break()
    doc.add_heading("근거 법령 및 참고 문서 목록", level=1)
    seen = {}
    for r in all_refs:
        name = r['source']
        seen[name] = seen.get(name, 0) + 1
    for name, cnt in seen.items():
        doc.add_paragraph(f"- {name} (참조 {cnt}회)")


def render_doc(project_name, depart_name, project_no, period, budget):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    doc.add_heading("연구개발계획서 자동생성 결과", 0)
    doc.add_paragraph(f"세부사업명: {depart_name}")
    doc.add_paragraph(f"과제명: {project_name}")
    doc.add_paragraph(f"과제번호: {project_no}")
    doc.add_paragraph(f"기간: {period}")
    doc.add_paragraph(f"예산: {budget} 천원\n")

    prompts = [build_prompt(s, project_name, depart_name, project_no, period, budget) for s in sections]
    generated = generate_text_batch(prompts)

    all_refs_flat = []
    for sec, gen_text in zip(sections, generated):
        doc.add_heading(sec['section'], level=1)
        cleaned = clean_generated_text(gen_text)
        refs = search_reference(sec["query"], topk=TOPK)
        validation = validate_output(sec, cleaned, refs)
        doc.add_paragraph(cleaned)
        p = doc.add_paragraph()
        p.add_run("[Eval_Result]").bold = True
        for v in validation:
            doc.add_paragraph(v)
        all_refs_flat.extend(refs)

    add_reference_section(doc, all_refs_flat)
    outpath = "RND_Report.docx"
    doc.save(outpath)
    print(f"[DONE] → {outpath}")


# ================================================================
# 9. 실행 예시
# ================================================================
if __name__ == "__main__":
    project_name = "AI 기반 흉부 X-ray 영상 판독 자동화 시스템 개발"
    depart_name  = "산업기술혁신사업"
    project_no   = "RS-2025-00123456"
    period       = "2025년 4월 28일 ~"
    budget       = "5,000,000"
    print(f"[INFO] 시작: {project_name} ({depart_name})")
    render_doc(project_name, depart_name, project_no, period, budget)
