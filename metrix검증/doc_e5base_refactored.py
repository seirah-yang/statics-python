# -*- coding: utf-8 -*-
"""
Doc Evaluator (Refactored, Combined Report)
- 지정된 '필수 제목 리스트'에 맞춰 각 문서의 섹션을 매칭해 평가
- 문서가 여러 개인 경우에도 통합 보고서 1개로 저장
- 기존 규칙 기반 KPI 추출, 간이 NLI/QA, 포맷·응집성 등 지표 재사용
"""

import os, re, json, math, unicodedata
from typing import List, Dict, Any, Tuple
from collections import namedtuple, Counter, defaultdict

# ------------------------
# 의존성
# ------------------------
try:
    import docx
except Exception as e:
    raise RuntimeError("python-docx가 필요합니다. `pip install python-docx` 후 재시작하세요.") from e

_EMBED_OK = False
_BM25_OK = False
try:
    from sentence_transformers import SentenceTransformer
    import numpy as np
    _EMBED_OK = True
except Exception:
    _EMBED_OK = False

try:
    from rank_bm25 import BM25Okapi
    _BM25_OK = True
except Exception:
    _BM25_OK = False

# ------------------------
# 유틸
# ------------------------
def _norm(s: str) -> str:
    # 공백·문장부호 축약 + 한글 정규화
    s = unicodedata.normalize("NFKC", (s or "").strip())
    s = re.sub(r"\s+", " ", s)
    return s

def split_ko_sentences(text: str) -> List[str]:
    text = re.sub(r"\s+", " ", text or "")
    sents = re.split(r"(?<=[\.?!])\s+|(?<=다\.)\s+", text)
    return [s.strip() for s in sents if s and s.strip()]

KPI_PATTERNS = [
    r"\b\d{1,3}\s?%(\s?(감소|증가|유지))",
    r"(오류율|에러율)\s*\d{1,3}\s?%",
    r"(정확도|완전성|재현율|정밀도)\s*(\d{1,3}\s?%)",
    r"(응답시간|지연)\s*\d+(\.\d+)?\s?(ms|초)",
    r"(처리량|TPS|QPS)\s*\d+(\.\d+)?",
    r"(기간|마감|데드라인)\s*(\d+\s?(일|주|개월|월|분기|년))",
    r"(비용|원가)\s*\d+(,\d{3})*(\.\d+)?\s?(원|만원|억)",
]

def extract_numbers_units(text: str) -> List[Tuple[float, str]]:
    out = []
    for m in re.finditer(r"(\d+(?:[\.,]\d+)?)(\s?%|ms|초|일|주|개월|월|분기|년|원|만원|억)?", text or ""):
        num = m.group(1).replace(",", "")
        unit = (m.group(2) or "").strip()
        try:
            out.append((float(num), unit))
        except Exception:
            pass
    return out

def number_match_quality(a, b) -> Tuple[int, bool]:
    match = 0; conflict = False
    for ax, au in a:
        for bx, bu in b:
            if au and bu and au == bu:
                if bx == 0:
                    continue
                rerr = abs(ax - bx) / (abs(bx) + 1e-6)
                if rerr <= 0.1: match += 1
                elif rerr >= 0.5: conflict = True
    return match, conflict

def keyword_overlap(a: str, b: str) -> float:
    ta = set(re.findall(r"[가-힣A-Za-z0-9]+", (a or "").lower()))
    tb = set(re.findall(r"[가-힣A-Za-z0-9]+", (b or "").lower()))
    if not ta or not tb: return 0.0
    return len(ta & tb) / len(ta | tb)

def ngram_redundancy(sentences: List[str], n: int = 3) -> float:
    grams = []
    for s in sentences:
        toks = re.findall(r"[가-힣A-Za-z0-9]+", (s or "").lower())
        grams += list(zip(*[toks[i:] for i in range(n)]))
    if not grams: return 0.0
    c = Counter(grams)
    dup = sum(v-1 for v in c.values() if v>1)
    return dup / (len(grams) + 1e-6)

def simple_coherence(sentences: List[str]) -> float:
    if len(sentences) < 2: return 0.5
    scores = []
    for i in range(len(sentences)-1):
        scores.append(keyword_overlap(sentences[i], sentences[i+1]))
    return sum(scores)/len(scores)

def format_score(sections: List[Dict[str,str]], required_titles: List[str]) -> float:
    titles = [(s["title"] or "").strip() for s in sections if s.get("title")]
    hit = 0
    for req in required_titles:
        if any(req in (t or "") for t in titles):
            hit += 1
    return hit / max(1, len(required_titles))

# ------------------------
# 파서
# ------------------------
class DocParser:
    def parse(self, docx_path: str):
        try:
            d = docx.Document(docx_path)
        except Exception:
            return None

        paras = [p.text.strip() for p in d.paragraphs if p.text and p.text.strip()]
        sections = []
        cur_title, cur_buf = None, []
        for p in d.paragraphs:
            style = getattr(p.style, "name", "") or ""
            text = (p.text or "").strip()
            if not text:
                continue
            if style.startswith("Heading") or "제목" in style:
                if cur_title is not None or cur_buf:
                    sections.append({"title": _norm(cur_title), "text": _norm("\n".join(cur_buf))})
                cur_title, cur_buf = text, []
            else:
                cur_buf.append(text)
        if cur_title is not None or cur_buf:
            sections.append({"title": _norm(cur_title), "text": _norm("\n".join(cur_buf))})

        full_text = _norm("\n".join(paras))
        sentences = split_ko_sentences(full_text)
        Doc = namedtuple("Doc", ["sections", "paragraphs", "sentences", "text"])
        return Doc(sections=sections, paragraphs=paras, sentences=sentences, text=full_text)

# ------------------------
# 검색기
# ------------------------
class HybridRetriever:
    def __init__(self, cfg: Dict[str, Any], corpus_texts: List[str]):
        self.cfg = cfg
        self.corpus = corpus_texts or []
        self.use_embed = False
        self.use_bm25  = False

        if _BM25_OK and self.corpus:
            tokenized = [self.tokenize(x) for x in self.corpus]
            self.bm25 = BM25Okapi(tokenized)
            self.use_bm25 = True

        self.embed_dim = None
        if _EMBED_OK and self.corpus and cfg["models"].get("embed"):
            try:
                self.embed = SentenceTransformer(cfg["models"]["embed"])
                self.corpus_vec = self.embed.encode(self.corpus, normalize_embeddings=True)
                self.embed_dim = self.corpus_vec.shape[1]
                self.use_embed = True
            except Exception:
                self.use_embed = False

    @staticmethod
    def tokenize(x: str) -> List[str]:
        return re.findall(r"[가-힣A-Za-z0-9]+", (x or "").lower())

    def topk(self, query: str, k: int = 5) -> List[Tuple[str, float]]:
        cand: Dict[int, float] = {}

        if self.use_bm25:
            scores = self.bm25.get_scores(self.tokenize(query))
            for i, s in enumerate(scores):
                if s > 0: cand[i] = cand.get(i, 0.0) + float(s)

        if self.use_embed:
            qv = self.embed.encode([query], normalize_embeddings=True)[0]
            sims = (self.corpus_vec @ qv)
            for i, s in enumerate(sims):
                cand[i] = cand.get(i, 0.0) + float(s) * 100.0

        if not cand:
            for i, t in enumerate(self.corpus):
                if query[:20] and query[:20] in (t or ""):
                    cand[i] = 1.0

        ranked = sorted(cand.items(), key=lambda x: x[1], reverse=True)[:k]
        return [(self.corpus[i], score) for i, score in ranked]

# ------------------------
# 경량 NLI/QA
# ------------------------
class NLIModel:
    def __init__(self, cfg): self.cfg = cfg
    def predict(self, claim: str, evidence: str) -> Dict[str, float]:
        claim_nums = extract_numbers_units(claim)
        evid_nums  = extract_numbers_units(evidence)
        entail = 0.33; contra = 0.33; neutr = 0.34
        match_cnt, conflict = number_match_quality(claim_nums, evid_nums)
        if match_cnt > 0 and not conflict:  entail, contra, neutr = 0.7, 0.1, 0.2
        elif conflict:                       entail, contra, neutr = 0.1, 0.7, 0.2
        elif keyword_overlap(claim, evidence) > 0.4:
                                             entail, contra, neutr = 0.5, 0.1, 0.4
        return {"entail": entail, "contra": contra, "neutral": neutr, "max_p": max(entail, contra, neutr)}

class BooleanQA:
    def __init__(self, cfg): self.cfg = cfg
    def yesno(self, question: str, context: str) -> str:
        ko = keyword_overlap(question.lower(), context)
        cn, en = extract_numbers_units(question), extract_numbers_units(context)
        match_cnt, conflict = number_match_quality(cn, en)
        if conflict: return "no"
        if match_cnt >= 1 or ko > 0.35: return "yes"
        return "no"

# ------------------------
# 평가 공통
# ------------------------
def extract_claims_from_text(text: str) -> List[str]:
    sents = split_ko_sentences(text or "")
    hits = []
    for s in sents:
        if any(re.search(p, s) for p in KPI_PATTERNS):
            hits.append(s)
    # 상한선 & 중복 제거
    seen, uniq = set(), []
    for s in hits:
        if s not in seen:
            uniq.append(s); seen.add(s)
    return uniq[:10]

def _fluency(sents: List[str]) -> float:
    if not sents: return 0.5
    lens = [len(s) for s in sents]
    mean_len = sum(lens)/len(lens)
    punct = sum(ch in ".,;:?!~" for s in sents for ch in s) / (sum(lens)+1e-6)
    score = 0.5 + 0.5 * math.tanh((mean_len-25)/50) - 0.2*abs(punct-0.03)
    return max(0.0, min(1.0, score))

# ------------------------
# 섹션 매칭(제목 유사도)
# ------------------------
def title_similarity(a: str, b: str) -> float:
    a = _norm(a or ""); b = _norm(b or "")
    if not a or not b: return 0.0
    return keyword_overlap(a, b)  # 간단히 Jaccard 기반

def map_sections_to_required(sections: List[Dict[str,str]], required_titles: List[str]) -> Dict[str, Dict[str,str]]:
    """
    각 required_title에 대해 문서 내 가장 유사한 섹션을 1:1 매칭.
    임계치 미만(예: 0.25)이면 미존재로 처리.
    """
    mapping = {rt: {"title": None, "text": ""} for rt in required_titles}
    used = set()
    for rt in required_titles:
        best_i, best_sim = -1, 0.0
        for i, sec in enumerate(sections):
            if i in used:
                continue
            sim = title_similarity(rt, sec.get("title", ""))
            if sim > best_sim:
                best_i, best_sim = i, sim
        if best_i >= 0 and best_sim >= 0.25:
            mapping[rt] = {"title": sections[best_i].get("title"), "text": sections[best_i].get("text", "")}
            used.add(best_i)
    return mapping

# ------------------------
# 섹션 평가
# ------------------------
def evaluate_section(rt_title: str, sec_text: str, cfg) -> Dict[str, Any]:
    # 검색 코퍼스: 해당 섹션 문장
    sents = split_ko_sentences(sec_text)
    retr = HybridRetriever(cfg, sents if sents else [""])
    claims = extract_claims_from_text(sec_text)
    if not claims:
        claims = sents[:5]  # 주장 없으면 대표문장 대체
    # 간이 NLI/QA
    nli = NLIModel(cfg["models"].get("nli"))
    qa  = BooleanQA(cfg["models"].get("qna"))

    entail = contra = unknown = 0
    details = []
    for c in claims:
        ev = retr.topk(c, k=3)
        best = ev[0][0] if ev else ""
        nout = nli.predict(c, best)
        conf = nout["max_p"]
        if conf >= 0.65:
            verdict = max((("entailment", nout["entail"]), ("contradiction", nout["contra"]), ("neutral", nout["neutral"])), key=lambda x:x[1])[0]
        else:
            verdict = "entailment" if qa.yesno(f"Is the claim supported? {c}", best).startswith("y") else "contradiction"
            conf = max(0.65, conf)
        if verdict == "entailment": entail += 1
        elif verdict == "contradiction": contra += 1
        else: unknown += 1
        details.append({"claim": c, "evidence": (best or "")[:200], "verdict": verdict, "confidence": conf})

    tot = max(1, entail+contra+unknown)
    accuracy = entail / tot
    flu = _fluency(sents)
    coh = simple_coherence(sents)
    red = ngram_redundancy(sents, n=3)

    return {
        "required_title": rt_title,
        "exists": bool(sec_text.strip()),
        "accuracy": float(accuracy),
        "fluency": float(flu),
        "coherence": float(coh),
        "redundancy": float(red),
        "kpi_count": sum(1 for s in sents if any(re.search(p, s) for p in KPI_PATTERNS)),
        "length_chars": len(sec_text),
        "details": details
    }

# ------------------------
# 통합 실행기
# ------------------------
def run_combined_report(docx_paths: List[str], required_titles: List[str], cfg: Dict[str, Any], out_path: str):
    parser = DocParser()
    all_results = defaultdict(list)  # key: required_title -> [ {doc, metrics...}, ... ]
    doc_level = []                   # 문서별 전체 포맷/응집성 등

    for path in docx_paths:
        doc = parser.parse(path)
        if not doc:
            doc_level.append({"doc": path, "parse_ok": False})
            continue

        # 문서 수준 포맷 점수(필수 제목 커버리지)
        fmt = format_score(doc.sections, required_titles)
        coh = simple_coherence(doc.sentences)
        flu = _fluency(doc.sentences)
        red = ngram_redundancy(doc.sentences, n=3)
        doc_level.append({
            "doc": path, "parse_ok": True,
            "format": float(fmt), "coherence": float(coh), "fluency": float(flu), "redundancy": float(red)
        })

        # 제목 매핑 → 섹션별 평가
        mapped = map_sections_to_required(doc.sections, required_titles)
        for rt in required_titles:
            sec_text = mapped[rt].get("text", "")
            sec_res = evaluate_section(rt, sec_text, cfg)
            sec_res["doc"] = path
            sec_res["matched_title"] = mapped[rt].get("title")
            all_results[rt].append(sec_res)

    # 통합 스코어 집계(제목별 평균)
    title_summary = {}
    for rt, items in all_results.items():
        if not items:
            continue
        def avg(k):
            vals = [x[k] for x in items if isinstance(x.get(k), (int, float))]
            return sum(vals)/len(vals) if vals else 0.0
        title_summary[rt] = {
            "coverage": sum(1 for x in items if x["exists"]) / max(1, len(items)),
            "accuracy": avg("accuracy"),
            "fluency": avg("fluency"),
            "coherence": avg("coherence"),
            "redundancy": avg("redundancy"),
            "avg_kpi": avg("kpi_count"),
            "avg_len": avg("length_chars"),
        }

    # 최종 보고서 작성
    lines = []
    lines.append("# 통합 문서 평가 보고서\n\n")
    lines.append(f"- 총 문서 수: {len(docx_paths)}\n")
    lines.append(f"- 필수 제목 수: {len(required_titles)}\n\n")

    lines.append("## 1) 문서 수준 개요\n")
    for row in doc_level:
        if not row.get("parse_ok"):
            lines.append(f"- {row['doc']}: 파싱 실패\n")
            continue
        lines.append(f"- {row['doc']}: Format={row['format']:.2f}, Coherence={row['coherence']:.2f}, Fluency={row['fluency']:.2f}, Redundancy={row['redundancy']:.2f}\n")
    lines.append("\n")

    lines.append("## 2) 제목(섹션)별 요약(평균)\n")
    lines.append("| 제목 | 커버리지 | Accuracy | Fluency | Coherence | Redundancy(↓) | KPI개수 | 길이 |\n")
    lines.append("|---|---:|---:|---:|---:|---:|---:|---:|\n")
    for rt in required_titles:
        ts = title_summary.get(rt, None)
        if not ts:
            lines.append(f"| {rt} | 0.00 | 0.00 | 0.00 | 0.00 | 0.00 | 0.00 | 0 |\n")
        else:
            lines.append(f"| {rt} | {ts['coverage']:.2f} | {ts['accuracy']:.2f} | {ts['fluency']:.2f} | {ts['coherence']:.2f} | {ts['redundancy']:.2f} | {ts['avg_kpi']:.2f} | {int(ts['avg_len'])} |\n")
    lines.append("\n")

    lines.append("## 3) 상세(문서×제목)\n")
    for rt in required_titles:
        lines.append(f"### [{rt}]\n")
        items = all_results.get(rt, [])
        if not items:
            lines.append("- 해당 섹션 없음\n\n");
            continue
        for it in items:
            lines.append(f"- 문서: {it['doc']}\n")
            lines.append(f"  - 매칭된 제목: {it.get('matched_title')}\n")
            lines.append(f"  - 존재여부: {it['exists']} | Accuracy={it['accuracy']:.2f}, Fluency={it['fluency']:.2f}, Coherence={it['coherence']:.2f}, Redundancy={it['redundancy']:.2f}, KPI={it['kpi_count']}, 길이={it['length_chars']}\n")
            # 세부 클레임 상위 3개만
            for d in it["details"][:3]:
                ev = d["evidence"].replace("\n", " ")
                lines.append(f"    - 주장: {d['claim']}\n")
                lines.append(f"      · 판정: {d['verdict']} (신뢰도 {d['confidence']:.2f})\n")
                lines.append(f"      · 근거: {ev[:200]}...\n")
        lines.append("\n")

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("".join(lines))
    return out_path

# ------------------------
# 기본 설정 & 실행 예시
# ------------------------
REQUIRED_TITLES = [
    "연구개발 목표",
    "연구개발 내용",
    "연구개발성과 활용계획 및 기대효과",
    "연구기획과제의 개요",
    "연구개발과제의 배경",
    "연구개발과제의 필요성",
    "보안등급의 분류 및 해당 사유",
    "기술개발 핵심어(키워드)",
    "연차별 개발목표",
    "연차별 개발내용 및 범위",
    "추진방법 및 전략",
    "과제 성과의 활용방안",
    "신규사업 신설의 기대효과",
    "사회적 가치 창출 계획",
    "사회적 가치창출의 기대효과",
    "경제적 성과창출의 기대효과",
    "신규 인력 채용 계획 및 활용 방안",
]

DOCX_LIST = ['/content/e5base_계획서930.docx']
# 설정 객체 (임베딩 모델 이름이 필요합니다. snunlp/KR-SBERT-V40K는 흔히 사용되는 한국어 SBERT 모델 이름 예시)
CONFIG = {"models": {"embed": "snunlp/KR-SBERT-V40K", "nli": {}, "qna": {}}}

if not DOCX_LIST:
    print("[알림] DOCX_LIST에 문서 경로를 17개 넣어 실행하세요.")
else:
    # 🚨 run_combined_report 호출 추가
    out_path = "/content/result/e5l0930_REPORT.txt"
    run_combined_report(DOCX_LIST, REQUIRED_TITLES, CONFIG, out_path) 
    print(" 보고서 저장 완료:", out_path)
